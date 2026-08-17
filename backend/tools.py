import os
import uuid
from io import BytesIO
import requests
from duckduckgo_search import DDGS
import edge_tts
import asyncio
import pypdf
import docx

class KlaraTools:
    @staticmethod
    def web_search(query, max_results=5):
        """Busca en DuckDuckGo y devuelve resultados."""
        try:
            with DDGS() as ddgs:
                results = list(ddgs.text(query, max_results=max_results))
                return results if results else []
        except Exception as e:
            return [{"error": f"Error en la búsqueda: {str(e)}"}]

    @staticmethod
    async def text_to_speech(text, voice="es-MX-DaliaNeural", output_file=None):
        """Convierte texto a voz usando edge-tts generando nombres únicos."""
        if not output_file:
            output_file = f"voice_{uuid.uuid4().hex[:8]}.mp3"
            
        try:
            communicate = edge_tts.Communicate(text, voice)
            await communicate.save(output_file)
            return output_file
        except Exception as e:
            raise RuntimeError(f"Error al generar audio TTS: {e}")

    @staticmethod
    def pdf_to_text(pdf_bytes):
        """Extrae texto de un PDF procesando páginas vacías o escaneadas."""
        try:
            reader = pypdf.PdfReader(BytesIO(pdf_bytes))
            text = ""
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
            return text if text.strip() else "El PDF no contiene texto extraíble (podría ser un escaneo o imagen)."
        except Exception as e:
            return f"Error al procesar el archivo PDF: {str(e)}"

    @staticmethod
    def docx_to_text(docx_bytes):
        """Extrae texto de un documento Word."""
        try:
            doc = docx.Document(BytesIO(docx_bytes))
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            return text if text.strip() else "El documento Word está vacío."
        except Exception as e:
            return f"Error al procesar el documento Word: {str(e)}"

    @staticmethod
    def create_docx(text, filename=None):
        """Crea un documento Word con el texto dado y nombre único."""
        if not filename:
            filename = f"documento_{uuid.uuid4().hex[:8]}.docx"
            
        try:
            doc = docx.Document()
            for line in text.split("\n"):
                doc.add_paragraph(line)
            doc.save(filename)
            return filename
        except Exception as e:
            raise RuntimeError(f"Error al crear el documento Word: {e}")